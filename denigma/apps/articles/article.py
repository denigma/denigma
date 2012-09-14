"""Structured Articles.

Structured article seperates the context from the representation in a highly
structured intuitive way.
It enables writing an article in the object-orientated programming language
of choice (i.e. Python) in a native feeling as writing code.

References, tables and figures, or other in text elements are marked with []
around them.

For tables and figures either a title or a filename should be used.

A reference can by a pubmed id, a link, or a names with dates, basicely everthing:
PMID: [8807293]
Link: [http://www.uwaging.org/genesdb/gene.php?id=35]
Single Author: [Wuttke, 2012]
Two authors: [Wuttke & Durchholz, 2013] OR [Wuttke and Durchholz, 2013.
Multiple authors: [Wuttke et al. 2012]

Multiple references are seperated by semicolon ';':
Multiple mixed references: [8807293; Wuttke, 2012; Wuttke & Durchholz, 2013]

It will be even possible to cite a citation within another one:
[1,2,3 in Wuttke et al. 2012]

By writing a structural article you do not need to worry about the type of
reference.

A function will try to identify the right citation via consulting Medline and
other similiar resources.

A search engine will read the reference and try to figure out what is the right
citation.

An enhanced function

An optional references list can be defined which might guide look up of the
right citation. Just drop in a list of string or a single string with endline
seperation and it will be used for mapping your references.

If a reference is unimbiugous or unsufficient a warning will be printed.

Common foreign language words such as latein or gene names will be italizes automatically.
e.g. 

It allows to create papers, reviews, presentations webpages, as well as a
thesis in a breeth.

Everything is captured in objects, which just need to be structured.

Output will can printed to the console, in Word, as PDF, rst (restructured text)
or HTML format.

#Comments can be added to a structured article with the '#' previuosing it between any
#element of the article (e.g. right after a pragraph to comment on this paragraph.

Changes can be made directly in any element. All elements of two versions of the same
will be compared to each via sequence alignment and changes can be outputed or
highlighted.

It is also possible to just restructure the order of the elements in a given article,
which will be reqonized by the comparision algothim as each element will be mapped to
its most similiar one and its position.

Basic concept of restuctured text will be enabled such as emphasis *bold* and 'italic'
and so furth.

Allows the output of different citations styles such as author name(s) + year or numbered references.

LaTeX examples:
http://www.latex-project.org/intro.html

LaTeX + Python:
https://github.com/nickloman/latex-pubmed
"""
import os
import sys
import re

#from sa import *

try: 
   from denigma.publication import Bibliography, Reference # as Bib
except ImportError:
   from library import Bibliography, Reference

# Utility functions:
try:
    from scripts.c import isdigit
    from scripts.continity import continum
except ImportError:
    try:
        from utils import isdigit, continum
    except ImportError:
        print("Could not import utility functions.")

# Outputs:
try: 
    from files.word import Word
except ImportError: 
     print("Could not import Word.")
try: 
    from docx import *
except ImportError:
    print("Could not import docx.")


LATIN = ['cis-', 'trans-', 'omics', 'in vivo', 'in vitro', 'in silico', 'in situ', 'C. elegans']


def show_exception_and_exit(exec_type, exc_value, tb):
    import traceback
    traceback.print_exception(exc_type, exc_value, tb)
    raw_input("Press key to exit.")
    sys.exit(-1)
sys.excepthook = show_exception_and_exit


class Article(object):
    def __init__(self):
        self.title = ''
        self.authors = None
        self.abstract = ''
        self.keywords = []
        self.glossary = {}
        self.subtitles = [] # Depriciated, use sections!
        self.section = None
        self.sections = []
        self.paragraphs = []
        self.figures = Figures()
        self.tables = Tables()
        self.references = []
        self.structure = [] # Defines the overall structure of the article.
        self.bibliography = Bibliography()
        self.citations = Citations()
        self.content = None

    def referencing(self, numbered=False, brackets=False, rest=True):
        """This function goes through all the paragraph and replaces the
        references for printing purposes."""
        if numbered: # Numbered referencing style.
            self.citations.numbered = True
        if brackets:
            self.citations.brackets = True
        if rest:
           self.citations.rest = True
            
        mapping = {}
        text = ''

        def replace(finding):
            match = finding.group(0)
            print "Found match at:", match, match[1:-1], mapping #paragraph.references
            if "Figure" in match or "Table" in match:
                items = match.split('; ')
                for item in items:
                     print item                
            if match[1:-1] in mapping: #paragraph.references
                #print "match in finding"
                #print "match is:", match[1:-1]
                if "Figure" in match or "Table" in match:
                    print "mapping is", match, match[1:-1]

                    #print "Identified %s as match" % match, mapping, mapping[match[1:-1]]
                    citationstyle = "(%s)"
                else:
                    if not numbered:
                        citationstyle = "(%s)"
                    else:
                        citationstyle = "[%s]"
                return citationstyle % mapping[match[1:-1]] #paragraph.references
            else:
                return match

        for paragraph in self.paragraphs:
            #print paragraph, type(paragraph)
            # find []
            # use regex to replace all in place citations.
            pattern = "\[(.+?)\]"
            findings = re.findall(pattern, str(paragraph))
            for finding in findings:
                if finding.startswith('Figure'): #finding in self.figures or
                    items = finding.replace('; ', ';').split(';')
                    for item in items:
                        figurename = item.replace('Figure: ', '')
                        self.figures.order.append(figurename)
                        numberedfigure = "Figure %s" % (self.figures.order.index(figurename)+1)
                        paragraph.references[item] = numberedfigure
                        #self.figures[figurename].title = numberedfigure
                elif finding in self.tables or finding.startswith('Table'):
                    items = finding.replace('; ', ';').split(';')
                    for item in items:
                        tablename = item.replace('Table: ', '')
                        m = re.findall(' (\w)$', tablename) # http://stackoverflow.com/questions/2362471/match-start-and-end-of-file-in-python-with-regex
                        if m:
                            panel = m[0]
                            print "Panel =", panel
                            tablename = tablename[:-2]
                            print "Tabelname =", tablename
                        else:
                            panel = ''
                        self.tables.order.append(tablename)
                        try:
                            numberedtable = "Table %s" % (self.tables.order.index(tablename)+1)
                            print "Numbered table =", numberedtable
                            paragraph.references[item] = numberedtable+panel
                            print "Assigning %s to %s" % (item, numberedtable+panel)
                            #self.tables[tablename] = numberedtable
                        except Exception as e:
                            print "Failed to refer to table: ", e
                elif finding.startswith('http://'): continue # Is link
                else:
                    paragraph.references[finding] = Refs(numbered, self.citations)
                #print finding

                regex = '(\d,\d|\d-\d)' #http://stackoverflow.com/questions/8609597/python-regular-expressions-or 
                # This regex is insufficient as it false tries to map references like those: [BMJ, 326, 1297-1299 (2003); BMJ, 337, a399 (2008)]
                m = re.findall(regex, finding)
                regex = '(\w,\s\d)'
                am = re.findall(regex, finding)
                #passed = False
                #while not passed:
                if not am and (m or (isdigit(finding) and int(finding) < 500)):
                    print "Found m"
                    items = finding.split(',')
                    print items
                    for item in items[:]:
                        if "-" in item:
                            start, end = item.split('-')
                            #try: 
                            items.extend(xrange(int(start), int(end)+1))
                            #except: passed = True
                            #print items
                    for item in map(str, items):
                        #print item
                        if item in self.references:
                            #print "item in references", item, type(item)
                            r = self.references[item]
                            paragraph.references[finding].append(r)
                            mapping = paragraph.references
                            self.citations.add(r)
                            text = re.sub(pattern, replace, str(paragraph)) #http://stackoverflow.com/questions/3997525/python-replace-with-regex
                        else:
                            pass
                            #paragraph.references[finding].append(item)
                            #mapping = paragraph.references
                            #text = re.sub(pattern, replace, str(paragraph))
                    paragraph.text = text
                    #print text
                    #passed = True
                else:
                    #print "did not found m"
                    #print paragraph.references[finding]
                    items = finding.replace('; ', ';').split(';')
                    for item in items:
                        #print "item", item
                        if item in self.figures or item.startswith('Figure'):
                            #print "found figure again"
                            pass#continue #print "FOUND and IMAGE"
                        elif item in self.tables or item.startswith('Table'): pass#continue
                        elif item.startswith('http://'): continue # Is link
                        try:
                            #print "Trying to map i pmid"
                            id = int(item)
                            #print "id =", id, type(id)

                            r = self.bibliography.find(id, printing=False)
                            #print "Found reference:", r[0].ref()
                            self.citations.add(r[0])
                            #print "added citation to citations:", len(self.citations), self.citations
                            paragraph.references[finding].append(r[0])
                            #print "Appended reference to paragraph references."
                            mapping = paragraph.references
                            #print "Mapped paragaph"
                            #pattern = "\[.+?\]"
                            
                            text = re.sub(pattern, replace, str(paragraph))
                            #print "paragraph is now subsitututed:", text

                        except Exception as e:
                            print "Exception occured:", e
                            print item, "is not a pmid"
                            item = item.replace('.,', '.').replace(' and ', ' & ')
                            if item in self.references:
                                r = self.references[item]
                                print "found reference", r
                                paragraph.references[finding].append(r or item)
                                #print paragraph.references[finding] ##
                                mapping = paragraph.references
                                #paragraph.text =
                                self.citations.add(r)
                                text = re.sub(pattern, replace, str(paragraph))
                                
                                #print len(self.citations), self.citations
                            elif item.replace('Figure: ' , '') in self.figures:
                                #print "Figure need to be handled", item, paragraph.references[finding]
                                mapping = paragraph.references
                                text = re.sub(pattern, replace, str(paragraph))
                                #print item, "is not in references."
                            elif item.replace('Table: ', '') in self.tables or item.replace('Table: ', '')[:-1] in self.tables:
                                mapping = paragraph.references
                                text = re.sub(pattern, replace, str(paragraph))
                            else:
                                print "Ref did not match anything:", finding
                                if "Table: " in item: # I don't know why this has to be here?
                                    mapping = paragraph.references
                                    text = re.sub(pattern, replace, str(paragraph))
                                else:
                                    try:
                                        if len(item) > 5: # Too short will probaly match something weird.
                                            # Search direct against Entrez and retrieve pubmed id:
                                            ids = self.bibliography.search(term=item)
                                            if len(ids) == 1:
                                                r = self.bibliography.find(int(ids[0]), printing=False)
                                                #print r
                                                self.citations.add(r[0])
                                                #print r[0]
                                                paragraph.references[finding].append(r[0])
                                                mapping = paragraph.references
                                                #print paragraph.references
                                                text = re.sub(pattern, replace, str(paragraph))
                                               #print text
                                            else:
                                                 raise Exception
                                        else:
                                             raise Exception
                                    except:                                    
                                        try:
                                            # Assumed to be alredy a reference:
                                            paragraph.references[finding].append(item)
                                            text = re.sub(pattern, replace, str(paragraph))
                                            print finding, paragraph.references[finding]
                                        except Exception as e:
                                            print e, finding, "is no reference."
                                
                    paragraph.text = text or paragraph.text
                    for term in LATIN:
                       paragraph.text = paragraph.text.replace(" %s " % term, " *%s* " % term).replace("(%s)" % term, " (*%s*)" % term)
                        


                        # Check whether reference is in references
                        # Create a reference instance
                        # declare type of representation
                        # insert representation as requested.
                        # pmid?

    ##                    def replace(finding):
    ##                        match = finding.group(0)
    ##                        print "match:", match
    ##                        if match[1:-1] in mapping:
    ##                            print "match in mapping"
    ##                            return mapping[match[1:-1]]
    ##                        else:
    ##                            return match
                        
                    #r = Reference(Entrez.read(Entrez.esummary(db='pubmed', id=id)))
                    #print "fetching reference"
                    
                    #print r
                    #print type(r)
                    ##print r.ref()#, findings
                    #print "mapping reference"
                    #mapping[finding] = r.ref()
                    #print id
                    #print type(r[0].ref())
                    #print "ref()", r[0].ref(), r[0].ref(), r[0].ref()
                    #print mapping[finding]
  
##                    #except:
##                        # It is not a pmid:
##                        if  "et al." in item:
##                            name, year = item.split('et al.')
##                            name, year = name.strip(), year.strip()
##                            for reference in self.references:
##                                if name in reference and year in reference:
##                                    mapping[finding] = reference
##                                    print "mapping:", mapping
##                                    #print reference
##                                    print "substituitng"
##                                    new_paragraph = re.sub("\[(.+?)\]", replace, str(paragraph))#, len(re.sub("\[(.+?)\]", replace, str(paragraph)))
##                                    print new_paragraph
                      
    def glossaring(self):
        """Creates an automatic generated glossary."""
        for paragraph in self.paragraphs:
            findings = re.findall("[A-Z]{2,}", str(paragraph)) #More tha two capital letters.
            #print findings
            
            abbreviation_explained = re.findall("[a-z]?[A-Z]{2,}[a-z]? \(.*?\)", str(paragraph)) #"([a-z]{1})?[A-Z]{2,}([a-z]{1})? \(.*?\)"
            for abbreviation_explaination in abbreviation_explained:
                #print abbreviation_explaination
                abbreviation = abbreviation_explaination.split(' (')[0]
                explaination = abbreviation_explaination.split('(')[1].split(')')[0]
                self.glossary[abbreviation] = explaination
                
            explaination_abbreviated = re.findall("\(([a-z]?[A-Z]{2,}[a-z]?)\)", str(paragraph))
            #print explaination_abbreviated
            for explaination_abbreviation in explaination_abbreviated:
                number_of_words = len(explaination_abbreviation)
                if explaination_abbreviation[-1] == 's': # Plural
                    number_of_words -= 1
                # Form a regular expression to fetch the explaination in front the abbreviation:
                word = '[A-Z,a-z,0-9]+ '
                regex = '%s\([a-z]?[A-Z]{2,}[a-z]?\)' % (word * number_of_words)
                #print regex
                results = re.findall(regex, str(paragraph))
                for result in results:
                    explaination = result.split('(')[0]
                    abbreviation = result.split(' (')[1].split(')')[0]
                    self.glossary[abbreviation] = explaination
                    #print result

    def counting(self):
        """Counts the character lenght of each section."""
        self.count = 0
        for section in self.sections:
            section.count = 0
            for paragraph in section.paragraphs:
                section.count += paragraph.count
            for subsection in section.subsections:
                section.count += len(subsection.title.replace('**', '*').replace('*', ''))
                for paragraph in subsection.paragraphs:
                    subsection.count += paragraph.count
                    section.count += paragraph.count
            print section.title, section.count
                        

    def structuring(self, type='Console', emphasis='**', count=False, rest=True):
        figuresAndTables = False # Will be set to True if there is a Tables a Figures section and prevent the generation of automtaic table and figures section.
        if count: self.counting()
        structure = self.structure
        structure.append(emphasis+str(self.title)+emphasis)
        if self.abstract:
            structure.append('\n'+emphasis+'Abstract:'+emphasis)
            if count: structure[-1] += ' [%s]' % self.abstract.count
            structure.append('%s' % self.abstract)
        if self.keywords:
            structure.append('\n'+emphasis+'Keywords'+emphasis+': %s' % self.keywords)
        if self.content:
            structure.append('\n'+emphasis+'Content:'+emphasis)
            structure.append(str(self.content))
        if self.glossary:
            structure.append('\n'+emphasis+'Glossary:'+emphasis)
            if count: structure[-1] += ' [%s]' % len(self.glossary)
            abbreviations = [structure.append('%s = %s' % (k, v)) for k,v in self.glossary.items()] #.title()
        structure.append('')
        if self.sections:
            for section in self.sections:
                structure.append(emphasis+str(section)+emphasis)
                if count: structure[-1] += ' [%s]' % section.count
                if section.paragraphs:
                    for paragraph in section.paragraphs:
                        if not rest: structure.append('%s\n' % str(paragraph).replace('**\n', '.** '))
                if section.subsections:
                    for subsection in section.subsections:
                        structure.append(emphasis+str(subsection)+emphasis)
                        if count: structure[-1] += ' [%s]' % subsection.count
                        for paragraph in subsection.paragraphs:
                            if not rest: structure.append('%s\n' % str(paragraph).replace('**\n', '.** '))
                if section.title == "Figures & Tables":
                    if count:
                        structure[-1] = structure[-1][:-3]+' [%s] [%s]' % (len(self.figures), len(self.tables))
                        #structure[-1] += ' [%s] [%s]' % (len(self.figures), len(self.tables))

                    structure.append(str(self.figures) or '')
                    structure.append(str(self.tables) or '')
                    figuresAndTables = True
##                else:
##                    for paragraph in section.paragraphs:
##                        structure.append('%s\n' % paragraph)
        else:
            structure.append('\n\n'.join(map(str, self.paragraphs))) # Call referencing here.


        if rest: structure.append('\nReferences\n==========') # ReST
        else: structure.append('\n'+emphasis+'References'+emphasis)
        if count: structure[-1] += ' [%s]' % len(self.citations)
        #structure.append(str("\n".join(self.references)))
        structure.append(str(self.citations))

        if not figuresAndTables:
            structure.append('')
            if self.figures:
                structure.append('**Figures**')
                structure.append(str(self.figures) or '')
                structure.append('')
            if self.tables:
                structure.append('**Tables**')
                structure.append('')
                structure.append(str(self.tables) or '')

        return structure


    def _string(self):
        """Transforms structure into a single string."""
        if not self.structure: self.structuring()
        return '\n'.join(self.structure)
    string = property(_string)

    def __repr__(self):
        """Prints brief consentive summary of the article
        Number of sections, parapgraphs and total word count"""
        return string

    def __len__(self):
        """Character count."""
        return len(self.string)

    def meta(self):
        """Return meta-information such as  table of conent"""

    def __str__(self):
        return self.string

    def findCitation(self):
        """Proof of concept whether a pmid is sufficient."""

    def output(self):
        article.glossaring()
        article.referencing()
        print self

    def printIt(self):
        for section in self.sections:
            print section
            print
            for subsection in section.subsections:
                    print subsection
                    print
                    #print len(subsection.paragraphs)
                    for paragraph in subsection.paragraphs:
                            print paragraph
                            print
    def word(self, input=None, output=None, close=False, insert=False):
        """Creates a Word output document."""
        if input:
            word = Word(input)
        else:
            word = Word()
        if insert: word.insert(self.string)
        else: word.write(self.string)
        if output:
            word.save(output)
        else:
            word.save()
        if close:
            word.close()


    def wd(self, input=None, output=None, close=False, insert=False):
        """Creates a Word output document including emphasis.
        Note this is unfortunaly very slow."""
        if input:
            word = Word(input)
        else:
            word = Word()
        #word.visible()
            
        strings = self.string.split('\n')
        strings.reverse()
        for string in strings:
            print string
            if "*" in string:
                word.insertEmphasis(string+'\n')
            else:
                word.insert(string+'\n')
        if output:
            word.save(output)
        else:
            word.save()
        if close:
            word.close()

    def wordIt(self, input=None, output=None, close=False):
        """Creates a Word output document including emphasis.
        Note this is unfortunaly very slow."""
        if input:
            word = Word(input)
        else:
            word = Word()

        for figure in self.figures.values():
            figure.filename =  os.path.join(os.getcwd(), figure.filename).replace('\\', '/') #self.name[:-3]

        word.i(self.string)
        for figure in self.figures.values():
            try:
                word.insertImage(figure.filename)
            except Exception as e:
                print e
                pass

        for table in self.tables.values():
            if not table.raw:
                word.insertTable(table)
        
                #figure.filename =  os.path.join(os.getcwd(), self.name[:-3], figure.filename) 
                #word.insertImage(os.path.join(os.getcwd(), self.name[:-3], figure.filename)) #__file__
##        if output:
##            try: word.save(output)
##            except: pass
##        else:
##            try: word.save()
##            except: pass
##        if close:
##            word.close()


    def docX(self):
        """Produces a docx representation of the article."""

        # Default set of relationships - these are the minimum components of a document
        relationships = relationshiplist()        
        
        # Make a new document tree - this is the main part of Word document.
        document = newdocument()

        # This xpath Location is where most interesting content lives
        docbody = document.xpath('/w:document/w:body', namespaces=nsprefixes)[0]

        docbody.append(heading(str(self.title), 1))
        docbody.append(heading('Abstract', 2))
        docbody.append(paragraph(str(self.abstract)))
        for section in self.sections:
            docbody.append(heading(str(section.title), 1))
            if section.paragraphs:
                for para in section.paragraphs:
                    docbody.append(paragraph(str(para), jc='both'))
            if section.subsections:
                for subsection in section.subsections:
                    docbody.append(heading(str(subsection), 2))
                    for para in subsection.paragraphs:
                        docbody.append(paragraph(str(para), jc='both'))


##        docbody.append(heading("Tables", 1))
##        for title, t in self.tables.items():
##            docbody.append(paragraph(title, style='ListNumber')) # ListBullet
##            docbody.append(table(str(t).split('\n')))
##
        # Add an image:
        relationships, picpara = picture(relationships, 'image1.png', 'Figure 1')
        docbody.append(picpara)

        # Search and replace:
        print 'Searching for soemthing in a paragraph ...',
        if search(docbody, 'the awesomeness'): print 'found it!'
        else: print 'nope.'

        print 'Searching for something in a heading ...',
        if search(docbody, '200 lines'): print 'found it!'
        else: print 'nope.'

        # Add a pagebreak:
        docbody.append(pagebreak(type='page', orient='portrait'))

        docbody.append(heading('References', 1))
        docbody.append(paragraph(str(self.citations)))
        
        # Create the properteies, contenttypes, and other support files:
        coreprops = coreproperties(title='Human Frontier Science Promotion',
                                   subject='Full Application',
                                   creator='Daniel Wuttke',
                                   keywords=self.keywords)
        appprops = appproperties()
        contypes = contenttypes()
        webings = websettings()
        wordships = wordrelationships(relationships)

        # Save the document:
        savedocx(document, coreprops, appprops, contypes, webings, wordships,
                 'HFSP.docx')

    def show(self):
        """Opens the output of the article."""
        word = Word()
        word.show()
       
    def wordle(self):
        """Creates a wordle of a doucment"""
        from library.wc import wordCloud
        string = []
        for section in self.sections:
            string.append(str(section))
            for subsection in section.subsections:
                string.append(str(subsection))
        for paragraph in self.paragraphs:
            string.append(str(paragraph))
        string = " ".join(string)
        print "Wordle string:", string

        # Remove references:
        string = re.sub(r'\[.*?\]', '', string) #http://stackoverflow.com/questions/640001/how-can-i-remove-text-within-parentheses-with-a-regex
        
        wordCloud(string.replace('**', '').replace('*', '').replace(' et al.', ' ').replace('JPM', '').replace(',', '').replace(';', '').replace(':', '')) #self.string

    def contenting(self):
        """Builds a table of content for this article."""
        self.content = Content(self.sections)
        return self.content
            

class Title:
    def __init__(self, name=''):
        self.name = name
        article.title = self

    def __repr__(self):
        return self.name


class Abstract:
    def __init__(self, text=''):
        self.text = text
        article.abstract = self
        self.count = len(self.text.replace('**', '').replace('*', ''))

    def __repr__(self):
        return self.text


class Authors:
    def __init__(self, *args, **kwargs):
        if isinstance(args, str):
            self.names = args.split('; ')
        elif isinstance(args, (tuple, list)):
            self.names = args
        #Include affliations as dict.
        if 'affliation' in kwargs:
            self.affliation = kwargs['affliation']
        article.authors = self

    def __repr__(self):
        return '; '.join(self.names)


class Keys(list):
    """Keywords describing an article.
    Instances of this class take either a list
    or a string containing a list (seperated either by "; " or by ", ")"""
    def __init__(self, words):
        super(Keys, self).__init__()
        if isinstance(words, str):
            if '; ' in words:
                words = words.split('; ')
            elif ', ' in words:
                words = words.split(', ')   
        self.extend(words)
        article.keywords = self

    def __repr__(self):
        return '; '.join(self)

class Subtitle:
    """Depriciated, use section instead."""
    def __init__(self, text, level):
        self.text = text
        self.level = level
        article.subtitles.append(self)


class Section:
    """A section which marks a specific topic in an article.
    It can be just followed by paragraphs it should contain or
    by used in by its text parameters and passing a mutline string containing
    paragraphs seperated by two newline '\n\n' characters."""
    def __init__(self, title, level=None, text=''):
        self.title = title
        self.level = level or 1
        self.text = text
        self.paragraphs = []
        article.sections.append(self)
        article.section = self
        self.subsections = []
        self.count = 0
    def __repr__(self):
        return self.title

    def append(self, string):
        #if isinstance(self.text, list):
        self.paragraphs.append(string)
        #elif isinstance(self, str):
        self.text += string

    def __len__(self):
        return len(self.paragraphs)


class Subsection:
    """A subsection is level lower of the preceding section."""
    def __init__(self, title, level=None, text=''):
        #Section.__init__(self, args, kwargs)
        #super(Subsection, self).__init__(args, kwargs)
        self.title = title
        self.level = level or 2
        self.text = text
        self.paragraphs = []
        #if isinstance(article.section, Section):
        article.sections[-1].subsections.append(self)
        article.subsections = []
        #else: pass
        #article.section
        article.section = self
        self.count = 0
        
    def __repr__(self):
        return self.title
    
  
class Paragraph:
    def __init__(self, text='', font_size=11, section=None):
        self.text = text
        self.font_size = font_size
        self.section = section or article.section or None
        self.references = {} #  A list of lists which harbours reference articles.
        article.paragraphs.append(self) # edit to add.
        if article.section:
            article.section.paragraphs.append(self)
            #print "append paragraph", self.text[:10], "to", article.section

    def count(self):
        return len(self.text.replace('**', '').replace('*', ''))
    count = property(count)

    def __repr__(self):
        return self.text



class Collection(dict):#, list):TypeError: Error when calling the metaclass base multiple bases have instance lay-out conflict
    """A flexible collection which can act either as dict or as list."""
    def __init__(self, type='dict'):
        self.type = type
        if type == 'dict':
            dict.__init__(self)
        elif type == 'list':
            list.__init__(self)
            
    def add(self, item):
        """Add an item to the collection."""
        if self.type == 'dict':
            self[item.title] = item
        elif self.type == 'list':
            self.append(item)


class Figures(Collection):
    """All the figures of the article."""
    def __init__(self, type='dict'):
        Collection.__init__(self, type)
        self.order = [] # Rename order to list.
        self.list = self.order
    def __repr__(self):
        string = []
        for figure in self.order:
            if figure:
                #string.append('\n%s\n**Figure: %s**\n' % (self[figure].legend, self[figure].title))
                string.append('\n%s\n**Figure %s: %s.** %s\n' % (self[figure].filename,
                                                             self.order.index(figure)+1,
                                                             self[figure].title,
                                                             self[figure].legend))
        return "".join(string)


class Figure:
    """A figure which can be any image file."""
    def __init__(self, title='', ext='.png', legend=None, section=None, filename=None):
        self.title = title
        self.filename = filename or title + ext
        self.legend = legend or ''
        self.section = section or article.section or None
        article.figures.add(self)

    def __repr__(self):
        return '%s %s' % (self.title, self.filename)


class Tables(Collection):
    """All the tables of the article."""
    def __init__(self, type='dict'):
        Collection.__init__(self)
        self.order = []
    def __repr__(self):
        string = []
        for table in self.order:
            if table and table in self:

                #string.append('**Table: '+self[table].title+'**\n')
                string.append('**Table %s: %s.** %s\n' % (self.order.index(table)+1,
                                                      self[table].title,
                                                      self[table].legend or ''))
                if self[table].raw:
                    string.append(self[table].data)
                else:
                    string.append(self[table].unique)
                string.append('\n\n')
        return "".join(string)

        

class Table:
    """A table which can be created by tuples, lists, or dicts."""
    def __init__(self, title='', legend=None, header=None, data=None, section=None, raw=True):
        if ":\n" in title:
            self.title, self.data = title.split(':\n')
        else:   
            self.title = title
            self.data = data.replace('    ', '\t')
        self.header = header
        self.legend = legend
        self.section = section or article.section or None
        self.raw = raw
        self.unique = '#'+self.title+'#'
        article.tables.add(self)

    def __repr__(self):
        return '%s %s %s' % (self.title, self.header, self.data)


class Content(list):
    """A table of content which is generated automatically."""
    def __init__(self, content):
        super(Content, self).__init__()
        self.add(content)
        self.paragraphTitles = False
    def add(self, item):
        if isinstance(item, list):
            self.extend(item)
        else:
            self.append(item)
    def __repr__(self):
        string = []
        for section in self:
            string.append(section.title)
            for subsection in section.subsections:
                string.append('- '+subsection.title)
                if self.paragraphTitles:
                    for paragraph in subsection.paragraphs:
                        if str(paragraph).startswith("**"):
                            string.append(' - '+str(paragraph).split('\n')[0].replace('**', ''))
        return '\n'.join(string)

            
class Glossary(dict):
    """A glossary of abbreviations and their explainations,
    automatically built from text.
    It is also possible to define terms seperatly and drop them in either,
    as list of tuples, dict or string with end of line chareacters as seperators."""
    def __init__(self, data):
        super(Glossary, self).__init__()
        self.representation = ''
        #abbreviations = []
        terms = data.split('\n')
        for term in terms:
            if not term: continue
            if ' = ' in term:
                abbreviation, explaination = term.split(' = ')
            elif ': ' in term:
                abbreviation, explaination = term.split(': ')
            self[abbreviation] = explaination
        
    def __repr__(self):
        representation.append('Glossary:')
        for k, v in self.items():
            representation.append('%s = %s' % (k,v))
        return '\n'.join(representation)


class Reviewers():
    """Mail reviewers to suggest."""
    def __init__(*args, **kwargs):
        pass


class Refs(list):
    """A in group of references within a paragraph."""
    def __init__(self, numbered=False, citations=None):
        list.__init__(self)
        self.numbered = numbered
        self.citations = citations # To deterimine the number in the reference citations.
        
    def __repr__(self):
        if not self.numbered:
            # If it was not identfied use the initial ref:
            return "; ".join([i.ref() if not isinstance(i, str) else i for i in self])
        else:
            try:
                numbers = continum([self.citations.index(repr(i))+1 for i in self])
            except:
                #print self
                numbers = i

            return numbers


class References(dict):
    """Reference which should be first consulted."""
    def __init__(self, *citations, **kwarks):
        if len(citations) == 1 and isinstance(citations[0], (str, unicode)):
            self.citations = []
            print citations
            #if '\\n' in citations: citations = citations[0].split('\\n')
            citations = citations[0].split('\n')
            for citation in citations:
                print citation
                if citation:
                    # Test whether it is numbered:
                    m = re.match('\s*(\d+)\.*\s*(.+)', citation)
                    if m:
                        m.group(2)
                        self[m.group(1)] = article.bibliography.find(m.group(2), printing=False)[0]
                        #print "m.group(1)", m.group(1), self[m.group(1)], m.group(2)
                    else:
                        try: r = article.bibliography.find(citation, printing=False)[0]
                        except: print "Print couldn't search for", citation
                        self.citations.append(r)
                        self[r.ref()] = r
        else:
            self.citations = citations
        article.references = self

    def __repr__(self):
        string = []
        for k in self.keys():
            string.append(k)
        return "\n".join(string)


class Citations(list):
    """Builds the article citation section."""
    def __init__(self):
        super(Citations, self).__init__()
        self.dictionary = {}
        if not isinstance(self, set):
            self.add = self.adding
        self.numbered = False
        self.brackets = False # [1] [2] [3]

        self.format = "{authors} ({year}) {title}. {journal} {volume}: {pages}." #author_last_name, author_intial, et al.
            
    def adding(self, item):
        if repr(item) not in self:
            item.number = len(self)
            self.append(repr(item))
            self.dictionary[repr(item)] = item
            
    def __str__(self):
        string = []

        if self.brackets: style = '[%s]\t%s'
        elif self.rest: style = '    %s.    %s'
        else: style = '%s.\t%s'
        
        if self.numbered:
            for number, i in enumerate(self):
                string.append(style % (number+1, self.dictionary[i].full()))
        else:
            for i in self:
                string.append(self.dictionary[i].full())
            
        return "\n".join(string)


class Acknowledgements:
    def __init__(self, text):
        self.text = text
        article.acknowledgements = self


class Summary(Abstract):
    """Summarizes the content of an article.
    It is similar to an abstract but designed for thesis."""
    def __init__(self, text):
        super(Summary, self).__init__(self)
        self.text = text
        article.summary = self


article = Article()


if __name__ == '__main__':
    Resources = 'D:/resources'
    Title("Title of an article which illustrates the concept of structured article.")
    Authors("D. A. Wuttke")
    Abstract("This is an *illustration* of a structured article, a **object**-orientated type of writing an article.")
    Keys("structured articles", "object-orientated", "better than latex")
    Subtitle("This is a subtitle for a section.", 1)
    Section("This is a Section")
    Paragraph('''Ok here are some number references [1,2-3].''')
    Paragraph("""This is the first paragraph [Interesting figure] and a reference with a semicolon [Ayyadevara et al., 2007].""")
    Paragraph("""This is the second paragraph containing [A nice table] with a reference [de Magalhaes, 2009].""")
    Paragraph('''This is the third paragraph illustrating [8807293; Ayyadevara et al., 2007] the identification of a pmid reference [8807293].''')
    Paragraph('''This paragraph shows the IA (Identification of abbreviation) for structured articles (SAs).
    Even induced pluripotent stem cells (iPSCs) can be found [4].''')

    Paragraph("This is a paragraph which belongs to the above section.")
    Paragraph("Yet another paragraph that belongs to this section.")
    Subsection("This is a subsection")
    Paragraph('''This is a paragraph belonging to the above subsection.''')
    Paragraph('''Some **bold** and some (*italic*)''')
    Paragraph('''**Natural Language Processing**\nComplementary to the above described exploitation of available database resources; text mining will be applied to extract data from the literature. Advanced techniques of natural langue processing will be used to identify relationships and entities annotations from the MedLine. For this step we will use the Natural Language Tool Kit (NLTK). The extracted information will be converted to the same schema as those from external databases but treated separately with greater caution. As automatically extracted literature information is prone to high rates of false information an additional step need to be pursued.''')

    Figure(title="Interesting figure", ext='.png')
    Table(title="A nice table",
          header=["Term", "p-value"],
          data=['structred', 0.01,
                'article', 0.004])
    References("""
de Magalhaes, J.P. (2009). Ageing research in the post-genome era: new technologies for an old problem. SEB experimental biology series 62, 99-115.
Ayyadevara, S., Alla, R., Thaden, J.J., and Shmookler Reis, R.J. (2008). Remarkable longevity and stress resistance of nematode PI3K-null mutants. Aging cell 7, 13-22.
##	1. 	Hansen M, Taubert S, Crawford D, Libina N, Lee SJ, Kenyon C (2007) Lifespan extension by conditions that inhibit translation in Caenorhabditis elegans. Aging Cell 6: 95-110.
##	2. 	Henis-Korenblit S, Zhang P, Hansen M, McCormick M, Lee SJ, Cary M, Kenyon C (2010) Insulin/IGF-1 signaling mutants reprogram ER stress response regulators to promote longevity. Proc Natl Acad Sci U S A 107: 9730-9735.
##	3. 	David DC, Ollikainen N, Trinidad JC, Cary MP, Burlingame AL, Kenyon C (2010) Widespread protein aggregation as an inherent part of aging in C. elegans. PLoS Biol 8: e1000450.

""")
#Wuttke and coworks. Structured Article, a Scientific Revolution. Structured Journal (2012).
#Wutte and Durchholz. Tada 2013.

    article.contenting()
    article.glossaring()
    article.referencing(numbered=True, brackets=True)
    print article
    #article.wordle()
    #article.wd(insert=True)
    article.wordIt()
    #article.docX()
else:
    pass
    #print article
    #sys.stdin.read(1)
    #os.system("pause")


#234567891123456789212345678931234567894123456789512345678961234567897123456789


